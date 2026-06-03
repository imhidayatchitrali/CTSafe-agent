from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.domain.models import DocxRenderInput, DocxRenderResult


class DocxRenderService:
    """Base editable DOCX renderer.

    The production milestone should replace this with the full CT Safe template
    integration. This base renderer intentionally avoids copying example company
    data from the reference DVR.
    """

    def render_draft(self, request: DocxRenderInput) -> DocxRenderResult:
        project_dir = request.output_dir / str(request.project.id)
        project_dir.mkdir(parents=True, exist_ok=True)
        output_path = project_dir / f"DVR_v{request.version}.docx"

        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)

        heading = document.add_heading("Documento di Valutazione dei Rischi", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("Bozza tecnica generata da CT Safe DVR Agent.")

        company = request.project.company
        table = document.add_table(rows=0, cols=2)
        for label, value in [
            ("Ragione sociale", company.company_name),
            ("Partita IVA", company.vat_number),
            ("Codice ATECO", company.ateco_code),
            ("Sede", company.site_address),
            ("Tipo documento", company.document_type),
            ("Categoria rischio", company.risk_category),
        ]:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value or "DATO MANCANTE"

        document.add_heading("Sezioni pilota", level=1)
        for section in request.sections:
            document.add_heading(f"{section.section_number} {section.title}", level=2)
            markdown = section.generated_markdown or "DATO MANCANTE"
            for paragraph in markdown.split("\n\n"):
                document.add_paragraph(paragraph.strip())

        document.add_heading("Note di controllo", level=1)
        document.add_paragraph(
            "Bozza editabile. Richiede revisione umana e integrazione con il "
            "template CT Safe completo prima della consegna finale."
        )

        document.save(output_path)

        warnings: list[str] = []
        if request.template_path and request.template_path.exists():
            warnings.append(
                "Reference CT Safe template detected but not copied to avoid leaking "
                "example company data in the thin slice."
            )
        else:
            warnings.append("Reference CT Safe template not found in this environment.")

        return DocxRenderResult(
            project_id=request.project.id,
            path=output_path,
            version=request.version,
            warnings=warnings,
        )

